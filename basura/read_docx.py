import docx

doc_path = r"C:\Users\GIRTEC\Claude\Projects\Jaguar MX\Cronograma_EvaluacionTecnica_JairMolina.docx"
doc = docx.Document(doc_path)

with open(r"C:\Users\GIRTEC\Claude\Projects\Jaguar MX\docx_output.txt", "w", encoding="utf-8") as f:
    f.write("--- PARAGRAPHS ---\n")
    for i, p in enumerate(doc.paragraphs):
        f.write(f"Para {i}: {p.text}\n")

    f.write("\n--- TABLES ---\n")
    for i, table in enumerate(doc.tables):
        f.write(f"Table {i}:\n")
        for row in table.rows:
            row_data = [cell.text.replace('\n', ' ') for cell in row.cells]
            f.write(" | ".join(row_data) + "\n")
