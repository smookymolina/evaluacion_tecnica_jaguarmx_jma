import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc_path = r"C:\Users\GIRTEC\Claude\Projects\Jaguar MX\Cronograma_EvaluacionTecnica_JairMolina.docx"
doc = docx.Document(doc_path)

# Update the checkboxes in the checklist table
for table in doc.tables:
    if len(table.rows) > 0 and "Entregable" in table.cell(0, 0).text:
        for row in table.rows[1:]:
            if '☐' in row.cells[3].text:
                row.cells[3].text = row.cells[3].text.replace('☐', '✅ COMPLETO')

# Find the indices of the Days to replace the content
day1_idx = -1
day2_idx = -1
day3_idx = -1
day4_idx = -1
decisiones_idx = -1

for i, p in enumerate(doc.paragraphs):
    if "DÍA 1 — Firmware y Arquitectura" in p.text:
        day1_idx = i
    elif "DÍA 2 — Simulación en Wokwi" in p.text:
        day2_idx = i
    elif "DÍA 3 — Esquemático y PCB en KiCad" in p.text:
        day3_idx = i
    elif "DÍA 4 — Documentación y Entrega Final" in p.text:
        day4_idx = i
    elif "5. Decisiones Técnicas Clave" in p.text:
        decisiones_idx = i

def clear_paragraphs(start_idx, end_idx):
    for i in range(start_idx + 1, end_idx):
        p = doc.paragraphs[i]
        p.text = ""
        p.style = doc.styles['Normal']

if all(idx != -1 for idx in [day1_idx, day2_idx, day3_idx, day4_idx, decisiones_idx]):
    # Clear old plan text
    clear_paragraphs(day1_idx, day2_idx)
    clear_paragraphs(day2_idx, day3_idx)
    clear_paragraphs(day3_idx, day4_idx)
    clear_paragraphs(day4_idx, decisiones_idx)
    
    # Insert new realized details
    # Day 1
    p = doc.paragraphs[day1_idx + 1]
    p.add_run("Fase ejecutada y completada con éxito.\n").bold = True
    p.add_run("• Arquitectura FSM: ").bold = True
    p.add_run("Se diseñó e implementó la Máquina de Estados Finitos en MicroPython con los estados INIT, READING, COOLING, IDLE y ERROR. Se implementó un flujo no bloqueante para mantener un escaneo térmico constante a 1 Hz.\n")
    p.add_run("• Adquisición Analógica: ").bold = True
    p.add_run("Se desarrolló el módulo adc_ntc.py aplicando la ecuación Beta (B=3435K validado con datasheet TEWA) y un buffer circular de 8 muestras con promedio móvil recortado (trimmed mean) para rechazo de ruido.\n")
    p.add_run("• Control de Potencia: ").bold = True
    p.add_run("Se implementó hbridge.py para manejar el actuador FIT0803 y fan.py para el MOSFET del extractor a 48V. Se integró un Watchdog Timer (WDT) de 8s, alimentándolo en ráfagas de 500ms durante el movimiento síncrono del actuador (3s) para garantizar un estado seguro sin falsos reinicios.\n")
    
    # Day 2
    p = doc.paragraphs[day2_idx + 1]
    p.add_run("Fase ejecutada y completada con éxito.\n").bold = True
    p.add_run("• Extensión de Plantilla: ").bold = True
    p.add_run("Se extendió la plantilla base en Wokwi utilizando un Raspberry Pi Pico y el chip personalizado TB6612FNG. Se mapearon los pines conforme al documento original (GP26/GP27 para sensores, GP0/6/7 para DIP, GP1/2/3/4 para salidas).\n")
    p.add_run("• Mejoras de Simulación: ").bold = True
    p.add_run("Se incorporó código no bloqueante usando millis() en C++ para evitar congelamiento durante la simulación de movimiento. Se agregó parpadeo dinámico del LED del ventilador (>60°C) para emular máxima potencia.\n")
    p.add_run("• Pruebas Interactivas: ").bold = True
    p.add_run("Se implementó una consola serial de comandos (SET:, TEMP:, EXT:) para alterar las condiciones térmicas en tiempo real, permitiendo probar la respuesta del sistema sin necesidad de mover potenciómetros físicos.\n")

    # Day 3
    p = doc.paragraphs[day3_idx + 1]
    p.add_run("Fase ejecutada y completada con éxito.\n").bold = True
    p.add_run("• Diseño de Esquemático: ").bold = True
    p.add_run("Se diseñó el circuito completo en KiCad v8, incluyendo el LDO AP2112K-3.3, los acondicionadores NTC con filtro RC (1kΩ + 100nF) y diodos de clamp BAT54S. Se corrigieron los pines cruzados (VENT y MOTOR_EN) identificados en la revisión de fase A.\n")
    p.add_run("• Ruteo de PCB Mixed-Signal: ").bold = True
    p.add_run("Se realizó un trazado en 2 capas dividiendo físicamente las tierras (AGND y PGND) y uniéndolas en un único punto (Net-Tie) para evitar ruido de conmutación. Las pistas de potencia del motor (48V y 5V) se trazaron a ≥1.5mm para soportar las corrientes de stall.\n")
    p.add_run("• Verificación: ").bold = True
    p.add_run("Ambos análisis (ERC y DRC) pasaron con cero errores en el diseño final.\n")

    # Day 4
    p = doc.paragraphs[day4_idx + 1]
    p.add_run("Fase ejecutada y completada con éxito.\n").bold = True
    p.add_run("• Documentación: ").bold = True
    p.add_run("Se respondieron exhaustivamente las 6 preguntas del cuestionario en el archivo ETISEJr_JairMolina.md, detallando las decisiones técnicas más difíciles (actuador bloqueante vs no bloqueante) y los errores detectados por la IA.\n")
    p.add_run("• Integración de IA: ").bold = True
    p.add_run("Se adjuntó el log de interacciones con Claude, demostrando cómo se usó para refactorizar la lógica del búfer circular ADC y la configuración del JSON en Wokwi.\n")
    p.add_run("• Empaquetado: ").bold = True
    p.add_run("Se desarrolló y ejecutó el script package_project.ps1 para compilar los entregables (Firmware, KiCad, MD) en el archivo ETISEJr_JairMolina.zip, cumpliendo estrictamente el formato de entrega solicitado.\n")

# Save changes
doc.save(doc_path)
print("Document updated successfully.")
